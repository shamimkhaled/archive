#!/usr/bin/env python3
"""Generate PDF and Word copies of the Sonali Bank Archive deployment spec."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
PDF_PATH = OUT_DIR / "Sonali-Bank-Archive-Deployment-Specification.pdf"
DOCX_PATH = OUT_DIR / "Sonali-Bank-Archive-Deployment-Specification.docx"

NAVY = colors.HexColor("#0F2C4C")
GREEN = colors.HexColor("#1A1A54")
ROW_ALT = colors.HexColor("#F3F6F4")
WARN_BG = colors.HexColor("#FFF6E8")
LINE = colors.HexColor("#D7DDD9")
MUTED = colors.HexColor("#4A5560")

FONT_REG = "/usr/share/fonts/noto/NotoSans-Regular.ttf"
FONT_MED = "/usr/share/fonts/noto/NotoSans-Medium.ttf"
FONT_BOLD = "/usr/share/fonts/noto/NotoSans-Bold.ttf"

WORKLOADS = {
    "pilot": {
        "label": "Pilot",
        "hint": "500 documents, ~20 users, light concurrent upload",
        "docs": "500",
        "users": "20",
        "concurrent": "5–15 typical",
        "app_ram": "4 GB",
        "postgres": "10 GB",
        "qdrant": "1 GB",
        "object": "50 GB",
        "egress": "50–100 GB/mo",
        "audit": "~220k rows over 3 years",
        "pdf_used": "3.3 GB",
        "pg_used": "0.5 GB",
        "qdrant_used": "0.18 GB",
    },
    "production": {
        "label": "Production",
        "hint": "5,000 documents, ~50 users, meeting-day spikes",
        "docs": "5,000",
        "users": "50",
        "concurrent": "15–30 on meeting days",
        "app_ram": "8 GB",
        "postgres": "20 GB",
        "qdrant": "4 GB",
        "object": "200 GB",
        "egress": "200–500 GB/mo",
        "audit": "~440k rows over 6 years",
        "pdf_used": "32.5 GB",
        "pg_used": "2 GB",
        "qdrant_used": "1.8 GB",
    },
    "growth": {
        "label": "Growth",
        "hint": "25,000 documents, ~80 users, concurrent OCR + heavy view",
        "docs": "25,000",
        "users": "80",
        "concurrent": "30+ with overlapping OCR",
        "app_ram": "8 GB (API) + 8 GB (worker)",
        "postgres": "50 GB",
        "qdrant": "16 GB",
        "object": "1 TB",
        "egress": "1 TB+/mo",
        "audit": "~440k hot + yearly object-store archive",
        "pdf_used": "162.5 GB",
        "pg_used": "8 GB",
        "qdrant_used": "9.1 GB",
    },
}

ENV_LABELS = {
    "local": "Local server / on-prem",
    "railway": "Railway",
    "aws": "Amazon Web Services",
    "digitalocean": "DigitalOcean",
}

# (headline, region, monthly, replica, callout_title, callout_body, services[])
# service = (name, sku, cpu, ram, disk, bandwidth, notes)
SPECS = {
    "local": {
        "pilot": (
            "Laptop or small office box running docker compose",
            "On-prem / LAN",
            "Hardware only",
            "1 process (uvicorn --reload)",
            "Local is for development and internal demo",
            "PDFs land in UPLOAD_DIR, there is no HA, and docker-compose.yml has no CPU/memory limits. Do not treat this as the durable Sonali Bank archive.",
            [
                ("Host machine", "Dev workstation", "4 cores min", "16 GB (8 GB floor)", "50 GB SSD + PDF volume", "1 Gbps LAN", "App 4 GB + Qdrant 2 GB + Postgres 1 GB + OS"),
                ("App (FastAPI)", "venv or Dockerfile", "shares host", "4 GB reserved", "UPLOAD_DIR local", "LAN only", "OCR + watermark in-process"),
                ("PostgreSQL 16", "compose postgres:16", "1 vCPU", "1 GB", "postgres_data volume", "localhost", "bcpdb; no backups unless you add them"),
                ("Qdrant 1.13", "compose qdrant/qdrant:v1.13.0", "1 vCPU", "2 GB", "qdrant_data volume", ":6333 / :6334", "1536-d cosine; summaries + chunks"),
                ("Redis 7", "compose redis:7-alpine", "shared", "256 MB", "none (ephemeral cache)", ":6379", "Search cache only — not a queue"),
                ("PDF storage", "local UPLOAD_DIR", "—", "—", "~4 GB used / 50 GB disk", "LAN", "No lifecycle or replication"),
            ],
        ),
        "production": (
            "On-prem server — only if the archive must stay inside the office",
            "On-prem / VPN",
            "Server + power",
            "1 uvicorn process",
            "Prefer object storage even on a local server",
            "A single disk holding PDFs is a single point of failure. Point STORAGE_BACKEND=s3 at MinIO or a NAS S3 gateway if this must stay on-prem.",
            [
                ("Host machine", "Tower / 1U server", "8 cores", "32 GB", "500 GB NVMe + backup", "1–10 Gbps LAN", "Leave 8 GB for OS and OCR spikes"),
                ("App (FastAPI)", "Docker image, port 8080", "2 vCPU", "8 GB", "container + /data/uploads", "LAN / reverse proxy", "Tesseract eng+ben, Poppler, 50 MB PDF cap"),
                ("PostgreSQL 16", "compose or native", "2 vCPU", "2 GB", "50 GB + nightly dump", "localhost", "7-day dump to a second disk"),
                ("Qdrant 1.13", "compose", "2 vCPU", "4 GB", "20 GB volume", "localhost", "~1.8 GB vectors at 5k docs"),
                ("Redis 7", "compose", "shared", "512 MB", "none", "localhost", "Fail-open cache, 1h TTL"),
                ("PDF storage", "MinIO or NAS", "—", "—", "200 GB usable", "LAN", "Do not keep production PDFs on the app disk"),
            ],
        ),
        "growth": (
            "Local/on-prem at 25k docs — split the worker or move to cloud",
            "On-prem only if policy requires",
            "Server fleet",
            "API + worker (do not run 2 schedulers)",
            "Do not scale compose replicas as-is",
            "In-process APScheduler would double meeting reminders, and the login rate limit is in-memory. Extract OCR/chunk indexing to a worker and keep one scheduler.",
            [
                ("API host", "App server", "4 cores", "8 GB", "100 GB OS", "10 Gbps LAN", "Thin API: search, view, meetings"),
                ("Ingest worker", "Same image, different command", "4 cores", "8 GB", "scratch for OCR", "LAN to MinIO", "Parse, 300 DPI OCR, embed, Qdrant upsert"),
                ("PostgreSQL 16", "dedicated VM", "4 vCPU", "8 GB", "100 GB + backups", "private LAN", "Partition audit_logs older than 3 years"),
                ("Qdrant", "dedicated VM", "4 vCPU", "16 GB", "80 GB SSD", "private LAN", "~9 GB HNSW at 25k docs; snapshot weekly"),
                ("Redis 7", "compose or native", "shared", "1 GB", "none", "private LAN", "Cache + future rate-limit / lock keys"),
                ("PDF storage", "MinIO erasure-coded", "—", "—", "1 TB usable", "LAN / 10 Gbps", "Versioning off; lifecycle optional"),
            ],
        ),
    },
    "railway": {
        "pilot": (
            "Railway app + Neon + Qdrant Cloud + Upstash + R2/S3",
            "Railway US/EU + nearest data plane",
            "~$40–90 / month",
            "1 Railway service",
            "Railway disk is ephemeral",
            "Set STORAGE_BACKEND=s3 (AWS S3 or Cloudflare R2). Local UPLOAD_DIR is lost on every deploy. Health check path is /healthz.",
            [
                ("App", "Railway Docker (this Dockerfile)", "2 vCPU", "4 GB", "ephemeral — not for PDFs", "shared ~100+ Mbps", "PORT from Railway; OCR needs the 4 GB"),
                ("PostgreSQL", "Neon 0.25 CU", "shared CU", "shared", "10 GB", "pooled ssl=require", "App rewrites postgresql:// to asyncpg"),
                ("Qdrant", "Qdrant Cloud 1 GB", "managed", "1 GB", "managed", "TLS :443", "QDRANT_URL + QDRANT_API_KEY"),
                ("Redis", "Upstash 256 MB", "managed", "256 MB", "none", "rediss://", "Or UPSTASH_REDIS_REST_* pair"),
                ("PDF objects", "R2 or S3 Standard", "—", "—", "50 GB provisioned", "egress billed", "~3 GB used at 500 docs"),
                ("LLM / embed", "OpenRouter or OpenAI", "—", "—", "—", "API calls", "text-embedding-3-small, 1536-d, batch 24"),
            ],
        ),
        "production": (
            "Documented production path in railway.toml",
            "Railway + ap-south-1 / R2 for objects if serving BD",
            "~$120–250 / month",
            "1 service (do not autoscale replicas yet)",
            "Stay on one app replica",
            "Reminders run on AsyncIOScheduler every 15 minutes inside the web process. A second replica would double Resend emails until that job is externalized.",
            [
                ("App", "Railway Pro 8 GB", "2 vCPU", "8 GB", "ephemeral", "shared; watch 502s on meeting days", "Watermark loads full PDF into RAM"),
                ("PostgreSQL", "Neon 1 CU", "1 CU", "managed", "20 GB", "pooled", "Point-in-time recovery on"),
                ("Qdrant", "Qdrant Cloud 4 GB", "managed", "4 GB", "managed", "TLS", "~1.8 GB vectors at 5k docs; headroom for HNSW"),
                ("Redis", "Upstash 256–512 MB", "managed", "512 MB", "none", "rediss://", "Cache stays small; SEARCH_CACHE_TTL=3600"),
                ("PDF objects", "S3 / R2", "—", "—", "200 GB", "200–500 GB/mo egress", "Each view ≈ 2× PDF size (fetch + watermarked body)"),
                ("Email", "Resend", "—", "—", "—", "low", "Invites + 48h/24h reminders"),
            ],
        ),
        "growth": (
            "Split ingest off the web dyno; grow Qdrant and object store",
            "Same as production",
            "~$400–800 / month",
            "API 1× + worker 1×",
            "Scale storage and Qdrant first",
            "Postgres metadata stays modest even at 25k docs. PDF bytes and HNSW RAM dominate cost. Add a worker before adding a second web replica.",
            [
                ("API", "Railway 8 GB", "2 vCPU", "8 GB", "ephemeral", "1 TB+ egress possible", "Search, view, meetings only"),
                ("Ingest worker", "Second Railway service", "2 vCPU", "8 GB", "ephemeral scratch", "to S3 + Qdrant + LLM", "Move _index_document_chunks + OCR here"),
                ("PostgreSQL", "Neon 2 CU", "2 CU", "managed", "50 GB", "pooled", "Yearly archive of audit_logs to R2"),
                ("Qdrant", "Qdrant Cloud 8–16 GB", "managed", "16 GB", "managed snapshots", "TLS", "~9 GB at 25k docs; consider quantization later"),
                ("Redis", "Upstash 1 GB", "managed", "1 GB", "none", "rediss://", "Use for rate-limit + reminder lock if going multi-replica"),
                ("PDF objects", "S3 / R2 + lifecycle", "—", "—", "1 TB", "1 TB+/mo", "Do not CDN raw docs — watermark is per-user, Cache-Control: no-store"),
            ],
        ),
    },
    "aws": {
        "pilot": (
            "ECS/Fargate or a single t3.medium + managed data plane",
            "ap-south-1 (Mumbai) or ap-southeast-1",
            "~$80–160 / month",
            "1 task",
            "Keep data in-region for Sonali Bank latency",
            "Run the app, RDS, ElastiCache, and S3 in ap-south-1. Qdrant Cloud has no Mumbai region — pick the nearest Qdrant site or run Qdrant on a t3.medium in the same VPC.",
            [
                ("App", "ECS Fargate or t3.medium", "2 vCPU", "4 GB", "20 GB EBS (OS only)", "Up to 5 Gbps burst", "ALB → container :8080, health /healthz"),
                ("PostgreSQL", "RDS db.t4g.small", "2 vCPU", "2 GB", "20 GB gp3", "VPC", "Single-AZ ok for pilot; ssl required"),
                ("Qdrant", "Qdrant Cloud 2 GB or t3.medium", "2 vCPU", "2 GB", "30 GB gp3", "VPC / TLS", "1536-d cosine collections"),
                ("Redis", "ElastiCache cache.t4g.micro", "2 vCPU", "0.5 GB", "none", "VPC", "Cache only"),
                ("PDF objects", "S3 Standard", "—", "—", "50 GB", "50–100 GB/mo egress", "Block public access; app uses get_object then watermarks"),
                ("Logs", "CloudWatch Logs", "—", "—", "90 days hot", "—", "Operational logs 3–6 months"),
            ],
        ),
        "production": (
            "t3.large API + RDS + S3 + Qdrant 4–8 GB",
            "ap-south-1",
            "~$200–450 / month",
            "1 task / instance",
            "Watermarked views double egress",
            "Every /view and /download fetches the object then returns a full stamped PDF. Budget 2× PDF size per request. CloudFront in front of raw S3 is the wrong pattern — docs are no-store and user-specific.",
            [
                ("App", "t3.large or Fargate 8 GB", "2 vCPU", "8 GB", "30 GB EBS", "Up to 5 Gbps", "Meeting-day headroom for concurrent watermarks"),
                ("PostgreSQL", "RDS db.t4g.medium", "2 vCPU", "4 GB", "50 GB gp3", "VPC", "7-day backup; Multi-AZ if board SLA requires it"),
                ("Qdrant", "Qdrant Cloud 4–8 GB or t3.large", "2 vCPU", "8 GB", "60 GB gp3", "VPC / TLS", "Snapshot to S3 weekly"),
                ("Redis", "cache.t4g.small", "2 vCPU", "1.4 GB", "none", "VPC", "Still oversized vs cache need — smallest HA option"),
                ("PDF objects", "S3 + Intelligent-Tiering", "—", "—", "200 GB", "200–500 GB/mo", "Server access logging on for 12 months"),
                ("Logs", "CloudWatch + S3 export", "—", "—", "90–180 days hot", "—", "Audit stays in RDS 3–6 years"),
            ],
        ),
        "growth": (
            "Split API and worker; RDS larger; Qdrant 16 GB",
            "ap-south-1",
            "~$600–1,200 / month",
            "API 1× + worker 1×",
            "Horizontal scale needs Redis locks first",
            "Before a second API task: move reminder sweep to EventBridge/cron with a Redis lock, and store login rate-limit counters in ElastiCache. Until then, only scale vertically.",
            [
                ("API", "Fargate 2 vCPU / 4 GB", "2 vCPU", "4 GB", "20 GB", "ALB", "No OCR on this task"),
                ("Ingest worker", "Fargate / t3.large 8 GB", "2 vCPU", "8 GB", "scratch", "to S3 + Qdrant", "OCR 300 DPI is RAM-heavy"),
                ("PostgreSQL", "RDS db.t4g.large", "2 vCPU", "8 GB", "100 GB gp3", "VPC", "Partition audit_logs; Multi-AZ"),
                ("Qdrant", "r6g.large or Qdrant 16 GB", "2 vCPU", "16 GB", "120 GB + snapshots", "VPC", "Quantize vectors if RAM > 70%"),
                ("Redis", "cache.t4g.small", "2 vCPU", "1.4 GB", "none", "VPC", "Cache + rate limit + scheduler lock"),
                ("PDF objects", "S3 Standard / Glacier IR mix", "—", "—", "1 TB", "1 TB+/mo", "CloudFront only if you cache watermarked bytes per user — usually not worth it"),
            ],
        ),
    },
    "digitalocean": {
        "pilot": (
            "App Platform 4 GB + managed Postgres/Redis + Spaces + Qdrant Cloud",
            "BLR1 (Bangalore) or SGP1",
            "~$50–100 / month",
            "1 App Platform instance",
            "Spaces speaks the S3 API this app already uses",
            "Set AWS_S3_ENDPOINT_URL to the Spaces endpoint, AWS_BUCKET_NAME to the Space, and STORAGE_BACKEND=s3. A single 8 GB droplet running compose is fine only behind VPN.",
            [
                ("App", "App Platform 4 GB (2 GB is demo-only)", "2 vCPU", "4 GB", "container ephemeral", "included tier + overage", "2 GB will OOM on 300 DPI OCR"),
                ("PostgreSQL", "Managed PG 1 vCPU", "1 vCPU", "1–2 GB", "10 GB", "VPC / public TLS", "Daily backups included"),
                ("Qdrant", "Qdrant Cloud 1 GB or 4 GB droplet", "2 vCPU", "1–4 GB", "25 GB volume", "private net", "Prefer Cloud over stuffing it on the app droplet"),
                ("Redis", "Managed Redis 1 GB", "shared", "1 GB", "none", "VPC", "Smallest managed size is already enough"),
                ("PDF objects", "Spaces 250 GB", "—", "—", "250 GB", "1 TB included typical", "~3 GB used at 500 docs"),
                ("Alt all-in-one", "Droplet 8 GB / 4 vCPU", "4 vCPU", "8 GB", "160 GB", "2–4 Gbps", "Compose stack; VPN only — not public prod"),
            ],
        ),
        "production": (
            "8 GB App Platform or dedicated droplet + managed data + Spaces",
            "BLR1 / SGP1",
            "~$140–280 / month",
            "1 instance",
            "Dedicated 8 GB for the app, not a shared 2 GB",
            "Watermarking and OCR share the same process. Meeting-day concurrent views need the extra RAM more than extra instances.",
            [
                ("App", "App Platform / droplet 8 GB", "4 vCPU", "8 GB", "OS only", "4–5 TB included typical", "Health check /healthz"),
                ("PostgreSQL", "Managed PG 2 vCPU", "2 vCPU", "4 GB", "30 GB", "VPC", "Standby optional"),
                ("Qdrant", "8 GB droplet + volume", "4 vCPU", "8 GB", "80 GB volume", "private net", "Or Qdrant Cloud 4–8 GB"),
                ("Redis", "Managed Redis 1 GB", "shared", "1 GB", "none", "VPC", "Cache only"),
                ("PDF objects", "Spaces 1 TB", "—", "—", "1 TB", "200–500 GB/mo extra possible", "CDN in front of Spaces does not help watermarked no-store PDFs"),
                ("Email", "Resend", "—", "—", "—", "low", "Same as other clouds"),
            ],
        ),
        "growth": (
            "API droplet + worker droplet + managed PG 4 vCPU + 16 GB Qdrant",
            "BLR1 / SGP1",
            "~$400–750 / month",
            "API 1× + worker 1×",
            "Grow Spaces and Qdrant; Postgres stays modest",
            "At 25k docs the object store is ~160 GB used (plan 1 TB). Qdrant wants 16 GB RAM. Postgres 60 GB is still mostly empty if you archive audit rows yearly.",
            [
                ("API", "Droplet 4 GB / 2 vCPU", "2 vCPU", "4 GB", "80 GB", "4–5 TB included", "No OCR"),
                ("Ingest worker", "Droplet 8 GB / 4 vCPU", "4 vCPU", "8 GB", "80 GB scratch", "private + Spaces", "OCR + embeddings"),
                ("PostgreSQL", "Managed PG 4 vCPU", "4 vCPU", "8 GB", "60 GB", "VPC", "Archive audit_logs > 3 years to Spaces"),
                ("Qdrant", "Droplet 16 GB + volume", "8 vCPU", "16 GB", "200 GB volume", "private net", "Snapshots to Spaces"),
                ("Redis", "Managed Redis 1 GB", "shared", "1 GB", "none", "VPC", "Add lock keys before a second API droplet"),
                ("PDF objects", "Spaces 1 TB+", "—", "—", "1 TB+", "1 TB+/mo possible", "Spaces CDN optional for static PWA assets only"),
            ],
        ),
    },
}

SCALE_TRIGGERS = [
    ("Archive > ~500 docs, or Railway/DO disk used for PDFs", "Move to S3 / R2 / Spaces immediately (STORAGE_BACKEND=s3)"),
    ("Qdrant RAM > 70% or search latency climbs", "Next Qdrant RAM tier (HNSW is RAM-bound). 1 → 4 → 8 → 16 GB"),
    ("Postgres disk > 70%", "Grow volume; partition/archive audit_logs older than 3 years"),
    ("Concurrent OCR uploads stall the UI", "Extract parse/OCR + _index_document_chunks to a worker; keep API thin"),
    ("Meeting-day view storms / 502s", "Vertical scale app 4 → 8 GB. Stay on 1 replica until scheduler is a cron"),
    ("Egress cost spikes", "Accept 2× PDF bytes per view. Do not CDN raw docs (watermark + no-store)"),
    ("Need 2+ app replicas", "Externalize reminders (cron + Redis lock) and login rate limit (Redis) first"),
    ("25k+ documents", "Qdrant 16 GB; consider quantized vectors. Postgres is still modest"),
]

LOG_STREAMS = [
    ("Audit trail", "Postgres audit_logs", "login, upload, view, download, access_request, user admin", "3–6 years hot", "~0.5–2 KB/row; <2 GB at 6 years"),
    ("App / uvicorn", "Railway logs, CloudWatch, journald", "OCR failures, index errors, 5xx", "3–6 months", "tens of MB/month"),
    ("Access / proxy", "ALB, Railway edge, Nginx", "HTTP method, status, latency", "3–6 months", "depends on QPS; drop after debug window"),
    ("Meeting notifications", "Postgres notification_events", "invites, 48h/24h reminders", "3–6 years (with meetings)", "small; meetings × invitees"),
    ("Object access", "S3 / R2 / Spaces server logs", "who fetched which key", "12 months", "low; enable on the bucket"),
    ("LLM / embedding usage", "OpenAI or OpenRouter dashboard", "tokens, spend, model", "12 months of invoices", "external — not in this DB"),
]


def _register_fonts() -> None:
    pdfmetrics.registerFont(TTFont("Noto", FONT_REG))
    pdfmetrics.registerFont(TTFont("Noto-Med", FONT_MED))
    pdfmetrics.registerFont(TTFont("Noto-Bold", FONT_BOLD))


def _styles():
    base = getSampleStyleSheet()
    styles = {
        "cover_kicker": ParagraphStyle(
            "cover_kicker",
            fontName="Noto-Med",
            fontSize=9,
            textColor=GREEN,
            tracking=1.2,
            spaceAfter=8,
        ),
        "cover_title": ParagraphStyle(
            "cover_title",
            fontName="Noto-Bold",
            fontSize=22,
            leading=28,
            textColor=NAVY,
            spaceAfter=8,
        ),
        "cover_sub": ParagraphStyle(
            "cover_sub",
            fontName="Noto",
            fontSize=11,
            leading=16,
            textColor=MUTED,
            spaceAfter=6,
        ),
        "h1": ParagraphStyle(
            "h1",
            fontName="Noto-Bold",
            fontSize=14,
            leading=18,
            textColor=NAVY,
            spaceBefore=14,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "h2",
            fontName="Noto-Bold",
            fontSize=11.5,
            leading=15,
            textColor=GREEN,
            spaceBefore=10,
            spaceAfter=5,
        ),
        "h3": ParagraphStyle(
            "h3",
            fontName="Noto-Med",
            fontSize=10,
            leading=13,
            textColor=NAVY,
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "body",
            fontName="Noto",
            fontSize=9,
            leading=13,
            textColor=colors.HexColor("#1A1F24"),
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "bullet",
            fontName="Noto",
            fontSize=9,
            leading=13,
            leftIndent=12,
            textColor=colors.HexColor("#1A1F24"),
            spaceAfter=3,
        ),
        "cell": ParagraphStyle(
            "cell",
            fontName="Noto",
            fontSize=7.5,
            leading=10,
            textColor=colors.HexColor("#1A1F24"),
        ),
        "cell_h": ParagraphStyle(
            "cell_h",
            fontName="Noto-Bold",
            fontSize=7.5,
            leading=10,
            textColor=colors.white,
        ),
        "note": ParagraphStyle(
            "note",
            fontName="Noto",
            fontSize=8,
            leading=11,
            textColor=MUTED,
            spaceAfter=6,
        ),
        "callout_title": ParagraphStyle(
            "callout_title",
            fontName="Noto-Bold",
            fontSize=9,
            leading=12,
            textColor=NAVY,
            spaceAfter=2,
        ),
        "callout_body": ParagraphStyle(
            "callout_body",
            fontName="Noto",
            fontSize=8.5,
            leading=12,
            textColor=colors.HexColor("#1A1F24"),
        ),
        "footer": ParagraphStyle(
            "footer",
            fontName="Noto",
            fontSize=7.5,
            textColor=MUTED,
            alignment=TA_CENTER,
        ),
        "meta": ParagraphStyle(
            "meta",
            parent=base["Normal"],
            fontName="Noto",
            fontSize=8.5,
            leading=12,
            textColor=MUTED,
        ),
    }
    return styles


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("\n", "<br/>"), style)


def _callout(title: str, body: str, styles) -> Table:
    inner = [
        [_p(title, styles["callout_title"])],
        [_p(body, styles["callout_body"])],
    ]
    table = Table(inner, colWidths=[170 * mm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), WARN_BG),
                ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#E2C48A")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (0, 0), 7),
                ("BOTTOMPADDING", (0, -1), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


def _table(headers: list[str], rows: list[list[str]], styles, col_widths: list[float]) -> Table:
    head = [_p(h, styles["cell_h"]) for h in headers]
    body = [[_p(c, styles["cell"]) for c in row] for row in rows]
    data = [head, *body]
    table = Table(data, colWidths=col_widths, repeatRows=1)
    cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Noto-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.25, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("BACKGROUND", (0, 1), (-1, -1), colors.white),
    ]
    for i in range(1, len(data)):
        if i % 2 == 0:
            cmds.append(("BACKGROUND", (0, i), (-1, i), ROW_ALT))
    table.setStyle(TableStyle(cmds))
    return table


def _header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFillColor(NAVY)
    canvas.rect(0, A4[1] - 12 * mm, A4[0], 12 * mm, fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont("Noto-Med", 8)
    canvas.drawString(18 * mm, A4[1] - 8 * mm, "Sonali Bank Archive System  ·  Document Archive Intelligent System")
    canvas.drawRightString(A4[0] - 18 * mm, A4[1] - 8 * mm, "Deployment specification")
    canvas.setFillColor(GREEN)
    canvas.rect(0, 0, A4[0], 10 * mm, fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont("Noto", 7.5)
    canvas.drawString(18 * mm, 4 * mm, "Confidential  ·  August 2026  ·  Derived from current architecture")
    canvas.drawRightString(A4[0] - 18 * mm, 4 * mm, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf_story(styles):
    story = []
    story.append(Spacer(1, 28 * mm))
    story.append(_p("SONALI BANK PLC", styles["cover_kicker"]))
    story.append(_p("Document Archive Intelligent System", styles["cover_title"]))
    story.append(_p("Deployment specification", styles["cover_title"]))
    story.append(
        _p(
            "Storage, RAM, database, vector database, bandwidth, history-log retention, "
            "and scale-up guidance for local, Railway, AWS, and DigitalOcean.",
            styles["cover_sub"],
        )
    )
    story.append(Spacer(1, 6 * mm))
    story.append(
        _p(
            "Version 1.0  ·  18 August 2026  ·  FastAPI + PostgreSQL 16 + Qdrant 1.13 + Redis 7 + S3-compatible object storage",
            styles["meta"],
        )
    )
    story.append(Spacer(1, 8 * mm))
    story.append(
        _table(
            ["Workload", "Documents", "Users", "App RAM", "Postgres", "Qdrant RAM", "Object store", "Egress"],
            [
                [
                    WORKLOADS[k]["label"],
                    WORKLOADS[k]["docs"],
                    WORKLOADS[k]["users"],
                    WORKLOADS[k]["app_ram"],
                    WORKLOADS[k]["postgres"],
                    WORKLOADS[k]["qdrant"],
                    WORKLOADS[k]["object"],
                    WORKLOADS[k]["egress"],
                ]
                for k in ("pilot", "production", "growth")
            ],
            styles,
            [22 * mm, 22 * mm, 18 * mm, 28 * mm, 20 * mm, 22 * mm, 22 * mm, 26 * mm],
        )
    )
    story.append(Spacer(1, 4 * mm))
    story.append(
        _p(
            "Recommended production starting point: 8 GB app RAM, 20 GB Postgres, 4 GB Qdrant, "
            "200 GB S3/R2/Spaces, one web replica, STORAGE_BACKEND=s3.",
            styles["note"],
        )
    )
    story.append(PageBreak())

    story.append(_p("1. What this system actually is", styles["h1"]))
    story.append(
        _p(
            "The Sonali Bank Archive System is a single FastAPI process (no Celery/RQ workers). "
            "OCR, LLM summarization, watermarking, and chunk indexing all share that process. "
            "The PWA is server-rendered Jinja2. Production PDFs must live in S3-compatible object storage "
            "because Railway (and typical PaaS) disks are ephemeral.",
            styles["body"],
        )
    )
    story.append(
        _p(
            "Hard constraints from the current codebase:",
            styles["body"],
        )
    )
    for item in [
        "PDFs only, 50 MB upload cap.",
        "OCR at 300 DPI (Tesseract eng+ben + Poppler in the Docker image) — this dominates RAM.",
        "View/download loads the full PDF twice (object fetch + watermarked response); no range streaming.",
        "Qdrant collections document_summaries + document_chunks, 1536-d cosine, 512-token chunks / 50 overlap.",
        "Redis is a fail-open search cache, not a job queue.",
        "One web replica only until the in-process APScheduler (15 min reminders) and in-memory login rate limit are externalized.",
    ]:
        story.append(_p(f"•  {item}", styles["bullet"]))

    story.append(_p("2. Sizing assumptions", styles["h1"]))
    story.append(
        _table(
            ["Input", "Assumption"],
            [
                ["Typical PDF", "5 MB (range 1–20 MB, hard cap 50 MB)"],
                ["Chunks per document", "~25 plus 1 summary vector"],
                ["Vector bytes", "1536 × 4 B ≈ 6 KB raw; ~12–16 KB with Qdrant HNSW"],
                ["Concurrent staff", "5–15 typical; 30+ on meeting days"],
                ["Users", "20–80 (admin / board_secretary / uploader / board_member)"],
                ["View bandwidth", "Each successful open ≈ 2× PDF size on the wire"],
            ],
            styles,
            [50 * mm, 130 * mm],
        )
    )
    story.append(Spacer(1, 3 * mm))
    story.append(
        _p(
            "Used storage (not provisioned SKU) grows mainly as PDF bytes. At 500 / 5,000 / 25,000 documents: "
            "PDF objects ~3.3 / 32.5 / 162.5 GB; Qdrant HNSW ~0.18 / 1.8 / 9.1 GB; Postgres ~0.5 / 2 / 8 GB "
            "(metadata + 6 years of audit at ~200 events/day). Postgres is not the cost driver.",
            styles["body"],
        )
    )

    story.append(_p("3. Workload tiers", styles["h1"]))
    story.append(
        _table(
            ["Metric", "Pilot", "Production", "Growth"],
            [
                ["Documents / users", "500 / 20", "5,000 / 50", "25,000 / 80"],
                ["Concurrency", WORKLOADS["pilot"]["concurrent"], WORKLOADS["production"]["concurrent"], WORKLOADS["growth"]["concurrent"]],
                ["App RAM target", "4 GB", "8 GB", "8 GB API + 8 GB worker"],
                ["Postgres provisioned", "10 GB", "20 GB", "50 GB"],
                ["Qdrant RAM", "1 GB", "4 GB", "16 GB"],
                ["Object storage", "50 GB", "200 GB", "1 TB"],
                ["Monthly egress", "50–100 GB", "200–500 GB", "1 TB+"],
                ["Audit history", "~220k / 3 yr", "~440k / 6 yr", "hot + yearly archive"],
            ],
            styles,
            [38 * mm, 47 * mm, 47 * mm, 48 * mm],
        )
    )

    story.append(_p("4. Environment specifications", styles["h1"]))
    story.append(
        _p(
            "Monthly dollar ranges are 2026 planning order-of-magnitude, not quotes. "
            "vCPU/RAM figures are reserved capacity, not average usage. Provisioned SKUs include "
            "headroom so OCR spikes and HNSW growth do not page.",
            styles["body"],
        )
    )

    for env_key in ("local", "railway", "aws", "digitalocean"):
        story.append(_p(f"4.{['local','railway','aws','digitalocean'].index(env_key)+1}  {ENV_LABELS[env_key]}", styles["h2"]))
        for wl in ("pilot", "production", "growth"):
            headline, region, monthly, replica, c_title, c_body, services = SPECS[env_key][wl]
            meta = WORKLOADS[wl]
            block = [
                _p(f"{meta['label']}  —  {headline}", styles["h3"]),
                _p(
                    f"Region: {region}  ·  Cost: {monthly}  ·  Replicas: {replica}  ·  "
                    f"{meta['hint']}",
                    styles["note"],
                ),
                _callout(c_title, c_body, styles),
                Spacer(1, 2.5 * mm),
                _table(
                    ["Service", "SKU", "vCPU", "RAM", "Disk", "Bandwidth", "Notes"],
                    [list(row) for row in services],
                    styles,
                    [24 * mm, 32 * mm, 18 * mm, 22 * mm, 28 * mm, 26 * mm, 30 * mm],
                ),
                Spacer(1, 4 * mm),
            ]
            story.append(KeepTogether(block))

    story.append(_p("5. History logs (3–6 months and 3–6 years)", styles["h1"]))
    story.append(
        _p(
            "Six streams. Keep board accountability in Postgres for years; drop noisy operational logs "
            "after a quarter or two. There is no TTL in code today — add a yearly archive job "
            "(CSV/Parquet to object storage) rather than deleting audit rows.",
            styles["body"],
        )
    )
    story.append(
        _table(
            ["Stream", "Store", "What is recorded", "Retain", "Volume"],
            [list(row) for row in LOG_STREAMS],
            styles,
            [32 * mm, 38 * mm, 48 * mm, 28 * mm, 34 * mm],
        )
    )
    story.append(Spacer(1, 2 * mm))
    story.append(
        _p(
            "At 200 audited actions/day: ~73k rows/year, ~220k at 3 years, ~440k at 6 years — well under 2 GB. "
            "Do not scale Postgres for logs. Scale it for backup RPO and connection pool. PDF bytes dominate cost.",
            styles["body"],
        )
    )
    story.append(
        _p(
            "Operational logs (app, proxy, meeting debug): 3–6 months hot. "
            "Audit trail and meeting notification events: 3–6 years. "
            "Object-storage access logs: 12 months. LLM usage: 12 months of invoices in the vendor dashboard.",
            styles["body"],
        )
    )

    story.append(_p("6. How to scale when traffic grows", styles["h1"]))
    story.append(
        _p(
            "Order of operations: object storage → Qdrant RAM → app RAM (4 → 8 GB) → ingest worker → "
            "replicas (only after Redis locks). Do not horizontally scale the current Dockerfile: "
            "a second web process would duplicate meeting reminder emails and split the in-memory login rate limit.",
            styles["body"],
        )
    )
    story.append(
        _table(
            ["Trigger", "Next step"],
            [list(row) for row in SCALE_TRIGGERS],
            styles,
            [78 * mm, 102 * mm],
        )
    )
    story.append(Spacer(1, 3 * mm))
    story.append(_p("Scale blockers in this codebase", styles["h2"]))
    story.append(
        _p(
            "•  Single replica — APScheduler reminder sweep and in-memory login rate limit (8 failures / 10 min) are process-local.",
            styles["bullet"],
        )
    )
    story.append(
        _p(
            "•  In-process OCR — parse_pdf + Tesseract 300 DPI + LLM summary run on the same uvicorn process as search and meetings.",
            styles["bullet"],
        )
    )
    story.append(
        _p(
            "•  Watermark bandwidth — view and download pull the full object, stamp it in memory, and return the whole PDF with Cache-Control: no-store.",
            styles["bullet"],
        )
    )

    story.append(_p("7. Production checklist", styles["h1"]))
    for item in [
        "APP_ENV=production, JWT_SECRET_KEY ≥ 32 characters, COOKIE_SECURE on.",
        "DATABASE_URL pointing at managed Postgres (Neon / RDS / DO Managed PG). App rewrites postgresql:// to asyncpg.",
        "QDRANT_URL + QDRANT_API_KEY (or self-hosted Qdrant with persistent volume).",
        "REDIS_URL=rediss://… or Upstash REST pair. Do not leave redis://127.0.0.1 on a PaaS.",
        "STORAGE_BACKEND=s3 with AWS_BUCKET_NAME, keys, region; set AWS_S3_ENDPOINT_URL for R2/Spaces/MinIO.",
        "OPENAI_API_KEY or OPENROUTER_API_KEY for summaries and text-embedding-3-small.",
        "Health check /healthz. One web replica until reminders and rate limits are externalized.",
        "Operational log retention 3–6 months; audit_logs retained 3–6 years with a yearly archive job.",
    ]:
        story.append(_p(f"•  {item}", styles["bullet"]))

    story.append(Spacer(1, 4 * mm))
    story.append(
        _p(
            "Sources: Dockerfile, docker-compose.yml, railway.toml, src/bcp_project/main_api.py, "
            "models.py, qdrant_store.py, chunker.py. Monthly costs are planning estimates only.",
            styles["note"],
        )
    )
    return story


def write_pdf() -> None:
    _register_fonts()
    styles = _styles()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
        title="Sonali Bank Archive — Deployment Specification",
        author="Sonali Bank Archive System",
        subject="Infrastructure sizing for local, Railway, AWS, and DigitalOcean",
    )
    doc.build(build_pdf_story(styles), onFirstPage=_header_footer, onLaterPages=_header_footer)


def _set_run_font(run, name="Calibri", size=11, bold=False, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x2C, 0x4C)


def _add_para(doc: Document, text: str, *, size=11, italic=False, space_after=8) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, color=RGBColor(0x1A, 0x1F, 0x24))
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)


def _shade_header(cell) -> None:
    tc = cell._tePr if False else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:fill"): "0F2C4C", qn("w:val"): "clear"})
    tcPr.append(shd)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(9)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _shade_header(cell)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Calibri"
    doc.add_paragraph()


def write_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    kicker = doc.add_paragraph()
    run = kicker.add_run("SONALI BANK PLC")
    _set_run_font(run, size=10, bold=True, color=RGBColor(0x1A, 0x1A, 0x54))

    title = doc.add_paragraph()
    run = title.add_run("Document Archive Intelligent System")
    _set_run_font(run, size=22, bold=True, color=RGBColor(0x0F, 0x2C, 0x4C))

    sub = doc.add_paragraph()
    run = sub.add_run("Deployment specification")
    _set_run_font(run, size=18, bold=True, color=RGBColor(0x0F, 0x2C, 0x4C))

    _add_para(
        doc,
        "Storage, RAM, database, vector database, bandwidth, history-log retention, "
        "and scale-up guidance for local, Railway, AWS, and DigitalOcean.",
        size=11,
    )
    _add_para(
        doc,
        "Version 1.0  ·  18 August 2026  ·  FastAPI + PostgreSQL 16 + Qdrant 1.13 + Redis 7 + S3-compatible object storage",
        size=10,
        italic=True,
    )

    _add_heading(doc, "Recommended production starting point", 1)
    _add_table(
        doc,
        ["Piece", "Target"],
        [
            ["App RAM", "8 GB (OCR + watermark share one process; 4 GB is a pilot floor)"],
            ["Postgres", "20–50 GB provisioned — metadata stays small"],
            ["Qdrant", "4 GB now → 16 GB at ~25k docs (HNSW is RAM-bound)"],
            ["PDFs", "S3 / R2 / Spaces — not Railway/app disk. ~5 MB × document count"],
            ["Bandwidth", "Each view ≈ 2× PDF size; plan 200–500 GB/mo at production"],
            ["Replicas", "One web instance until reminders and login rate-limit move off in-memory"],
            ["Logs", "Operational 3–6 months; audit 3–6 years (<2 GB even at 6 years)"],
        ],
    )

    _add_heading(doc, "1. What this system actually is", 1)
    _add_para(
        doc,
        "The Sonali Bank Archive System is a single FastAPI process (no Celery/RQ workers). "
        "OCR, LLM summarization, watermarking, and chunk indexing all share that process. "
        "The PWA is server-rendered Jinja2. Production PDFs must live in S3-compatible object storage "
        "because Railway (and typical PaaS) disks are ephemeral.",
    )
    for item in [
        "PDFs only, 50 MB upload cap.",
        "OCR at 300 DPI (Tesseract eng+ben + Poppler) — this dominates RAM.",
        "View/download loads the full PDF twice (object fetch + watermarked response); no range streaming.",
        "Qdrant collections document_summaries + document_chunks, 1536-d cosine, 512-token chunks / 50 overlap.",
        "Redis is a fail-open search cache, not a job queue.",
        "One web replica only until the in-process APScheduler (15 min reminders) and in-memory login rate limit are externalized.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _add_heading(doc, "2. Sizing assumptions", 1)
    _add_table(
        doc,
        ["Input", "Assumption"],
        [
            ["Typical PDF", "5 MB (range 1–20 MB, hard cap 50 MB)"],
            ["Chunks per document", "~25 plus 1 summary vector"],
            ["Vector bytes", "1536 × 4 B ≈ 6 KB raw; ~12–16 KB with Qdrant HNSW"],
            ["Concurrent staff", "5–15 typical; 30+ on meeting days"],
            ["Users", "20–80 (admin / board_secretary / uploader / board_member)"],
            ["View bandwidth", "Each successful open ≈ 2× PDF size on the wire"],
        ],
    )
    _add_para(
        doc,
        "Used storage (not provisioned SKU) grows mainly as PDF bytes. At 500 / 5,000 / 25,000 documents: "
        "PDF objects ~3.3 / 32.5 / 162.5 GB; Qdrant HNSW ~0.18 / 1.8 / 9.1 GB; Postgres ~0.5 / 2 / 8 GB "
        "(metadata + 6 years of audit at ~200 events/day). Postgres is not the cost driver.",
    )

    _add_heading(doc, "3. Workload tiers", 1)
    _add_table(
        doc,
        ["Metric", "Pilot", "Production", "Growth"],
        [
            ["Documents / users", "500 / 20", "5,000 / 50", "25,000 / 80"],
            ["Concurrency", WORKLOADS["pilot"]["concurrent"], WORKLOADS["production"]["concurrent"], WORKLOADS["growth"]["concurrent"]],
            ["App RAM target", "4 GB", "8 GB", "8 GB API + 8 GB worker"],
            ["Postgres provisioned", "10 GB", "20 GB", "50 GB"],
            ["Qdrant RAM", "1 GB", "4 GB", "16 GB"],
            ["Object storage", "50 GB", "200 GB", "1 TB"],
            ["Monthly egress", "50–100 GB", "200–500 GB", "1 TB+"],
            ["Audit history", "~220k / 3 yr", "~440k / 6 yr", "hot + yearly archive"],
        ],
    )

    _add_heading(doc, "4. Environment specifications", 1)
    _add_para(
        doc,
        "Monthly dollar ranges are 2026 planning order-of-magnitude, not quotes. "
        "vCPU/RAM figures are reserved capacity, not average usage.",
    )
    for env_key in ("local", "railway", "aws", "digitalocean"):
        _add_heading(doc, ENV_LABELS[env_key], 2)
        for wl in ("pilot", "production", "growth"):
            headline, region, monthly, replica, c_title, c_body, services = SPECS[env_key][wl]
            meta = WORKLOADS[wl]
            _add_heading(doc, f"{meta['label']} — {headline}", 3)
            _add_para(
                doc,
                f"Region: {region}  ·  Cost: {monthly}  ·  Replicas: {replica}  ·  {meta['hint']}",
                size=10,
                italic=True,
            )
            note = doc.add_paragraph()
            run = note.add_run(f"Note: {c_title}. {c_body}")
            _set_run_font(run, size=10, color=RGBColor(0x5A, 0x3A, 0x00))
            _add_table(
                doc,
                ["Service", "SKU", "vCPU", "RAM", "Disk", "Bandwidth", "Notes"],
                [list(row) for row in services],
            )

    _add_heading(doc, "5. History logs (3–6 months and 3–6 years)", 1)
    _add_para(
        doc,
        "Six streams. Keep board accountability in Postgres for years; drop noisy operational logs "
        "after a quarter or two. There is no TTL in code today — add a yearly archive job "
        "(CSV/Parquet to object storage) rather than deleting audit rows.",
    )
    _add_table(
        doc,
        ["Stream", "Store", "What is recorded", "Retain", "Volume"],
        [list(row) for row in LOG_STREAMS],
    )
    _add_para(
        doc,
        "At 200 audited actions/day: ~73k rows/year, ~220k at 3 years, ~440k at 6 years — well under 2 GB. "
        "Do not scale Postgres for logs. PDF bytes dominate cost.",
    )

    _add_heading(doc, "6. How to scale when traffic grows", 1)
    _add_para(
        doc,
        "Order of operations: object storage → Qdrant RAM → app RAM (4 → 8 GB) → ingest worker → "
        "replicas (only after Redis locks). Do not horizontally scale the current Dockerfile.",
    )
    _add_table(doc, ["Trigger", "Next step"], [list(row) for row in SCALE_TRIGGERS])
    _add_heading(doc, "Scale blockers in this codebase", 2)
    for item in [
        "Single replica — APScheduler reminder sweep and in-memory login rate limit (8 failures / 10 min) are process-local.",
        "In-process OCR — parse_pdf + Tesseract 300 DPI + LLM summary run on the same uvicorn process as search and meetings.",
        "Watermark bandwidth — view and download pull the full object, stamp it in memory, and return the whole PDF with Cache-Control: no-store.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _add_heading(doc, "7. Production checklist", 1)
    for item in [
        "APP_ENV=production, JWT_SECRET_KEY ≥ 32 characters, COOKIE_SECURE on.",
        "DATABASE_URL pointing at managed Postgres (Neon / RDS / DO Managed PG). App rewrites postgresql:// to asyncpg.",
        "QDRANT_URL + QDRANT_API_KEY (or self-hosted Qdrant with persistent volume).",
        "REDIS_URL=rediss://… or Upstash REST pair. Do not leave redis://127.0.0.1 on a PaaS.",
        "STORAGE_BACKEND=s3 with AWS_BUCKET_NAME, keys, region; set AWS_S3_ENDPOINT_URL for R2/Spaces/MinIO.",
        "OPENAI_API_KEY or OPENROUTER_API_KEY for summaries and text-embedding-3-small.",
        "Health check /healthz. One web replica until reminders and rate limits are externalized.",
        "Operational log retention 3–6 months; audit_logs retained 3–6 years with a yearly archive job.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _add_para(
        doc,
        "Sources: Dockerfile, docker-compose.yml, railway.toml, src/bcp_project/main_api.py, "
        "models.py, qdrant_store.py, chunker.py. Monthly costs are planning estimates only.",
        size=9,
        italic=True,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    write_pdf()
    write_docx()
    print(f"Wrote {PDF_PATH}")
    print(f"Wrote {DOCX_PATH}")
